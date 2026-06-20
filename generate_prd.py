# -*- coding: utf-8 -*-
"""
天猫数据分析项目 PRD (产品需求文档)
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from datetime import datetime

def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(doc, headers, rows, col_widths=None, header_color="667eea"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        run = cell.paragraphs[0].add_run(str(h))
        run.bold = True; run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)
    for i, row in enumerate(rows):
        bg = "F8F9FF" if i % 2 == 0 else "FFFFFF"
        for j, val in enumerate(row):
            cell = table.rows[i+1].cells[j]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
            if j == 0: run.bold = True
            set_cell_shading(cell, bg)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def add_body(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10); run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(10); run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

def build_prd():
    doc = Document()

    # === 封面 ===
    title = doc.add_heading('天猫用户销售数据分析平台\n产品需求文档 (PRD)', level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x66, 0x7e, 0xea); run.font.size = Pt(22)

    add_body(doc, f'版本: v2.0 | 日期: {datetime.now().strftime("%Y-%m-%d")} | 状态: 已交付')
    add_body(doc, '数据源: MySQL tmall_data (5表) | 技术栈: Python + pandas + sklearn + Plotly')
    doc.add_page_break()

    # === 1. 项目背景 ===
    add_heading(doc, '一、项目背景与目标', 1)
    add_body(doc, '天猫平台需要建立一套完整的用户销售数据分析体系，实现从"看数据"到"用数据"的跨越。'
              '当前数据库积累了用户、订单、产品、行为四类核心数据，但缺乏系统性的分析报表和用户分层能力。')
    add_body(doc, '核心目标:')
    add_bullet(doc, '搭建多维BI可视化报表，覆盖营收、用户、产品、行为、评价5大主题')
    add_bullet(doc, '建立用户生命周期分析体系（DAU/MAU/留存/流失/LTV）')
    add_bullet(doc, '使用机器学习完成用户聚类分群，支撑精细化运营')
    add_bullet(doc, '输出数据分析报告，为业务经营提供数据驱动的决策建议')

    # === 2. 项目价值 ===
    add_heading(doc, '二、项目期望收获与价值', 1)

    add_heading(doc, '2.1 业务价值', 2)
    add_bullet(doc, '收入增长: 通过转化率优化和用户唤醒，预计可提升月度营收5-15%')
    add_bullet(doc, '成本节约: 用户分群运营降低无效营销投放，预计节省营销成本20-30%')
    add_bullet(doc, '用户资产保值: 沉默流失用户召回率目标15-20%，挽回高价值用户资产')
    add_bullet(doc, '决策效率: 从"周报看数"升级为"实时看板+自动预警"，决策响应速度提升10倍')

    add_heading(doc, '2.2 技术价值', 2)
    add_bullet(doc, '建立可复用的数据分析Pipeline（数据提取→特征工程→模型训练→可视化→报告生成）')
    add_bullet(doc, '积累用户聚类模型（K-Means RFM分段），可作为推荐系统和定价模型的基础')
    add_bullet(doc, '沉淀数据指标体系（12+核心指标定义和计算公式），统一团队数据口径')
    add_bullet(doc, '形成"Python全栈数据分析"技术方案，可在其他业务线快速复制')

    add_heading(doc, '2.3 组织价值', 2)
    add_bullet(doc, '数据文化: 推动业务团队从经验决策转向数据驱动决策')
    add_bullet(doc, '技能沉淀: 项目的Python代码、SQL模板、分析框架可复用至其他项目')
    add_bullet(doc, '跨部门协同: BI报表可作为运营、产品、市场团队的统一数据看板')
    add_bullet(doc, 'AI能力验证: 验证AI辅助数据分析的可行性和效率，为后续AI应用提供经验')

    # === 3. 数据源 ===
    add_heading(doc, '三、数据源与表结构', 1)
    add_body(doc, '数据库: MySQL tmall_data (localhost, utf8mb4)')

    add_heading(doc, '3.1 users — 用户基础信息表 (5,000行)', 2)
    add_styled_table(doc,
        ['字段', '类型', '说明'],
        [['user_id', 'VARCHAR(255)', '用户唯一标识'],
         ['age', 'INT', '年龄 (18-63)'],
         ['gender', 'VARCHAR(255)', '性别 (男/女)'],
         ['province', 'VARCHAR(255)', '省份'],
         ['city', 'VARCHAR(255)', '城市'],
         ['registration_date', 'VARCHAR(255)', '注册时间'],
         ['member_level', 'VARCHAR(255)', '会员等级 (普通/铜牌/银牌/金牌/钻石)'],
         ['account_balance', 'DOUBLE', '账户余额'],
         ['credit_score', 'INT', '信用分 (400-800)']],
        col_widths=[4, 4, 8]
    )

    add_heading(doc, '3.2 orders — 订单表 (15,000行)', 2)
    add_styled_table(doc,
        ['字段', '类型', '说明'],
        [['order_id', 'VARCHAR(255)', '订单唯一标识'],
         ['user_id', 'VARCHAR(255)', '用户ID (外键→users)'],
         ['product_id', 'VARCHAR(255)', '产品ID (外键→products)'],
         ['quantity', 'INT', '购买数量'],
         ['order_date / order_date_date', 'VARCHAR / DATE', '下单时间'],
         ['order_status', 'VARCHAR(255)', '订单状态 (待付款/已付款/已发货/已收货/已完成/已取消/已退款)'],
         ['payment_method', 'VARCHAR(255)', '支付方式 (支付宝/微信/银行卡/信用卡/花呗)'],
         ['unit_price', 'DOUBLE', '单价'],
         ['total_amount', 'DOUBLE', '订单标价总额'],
         ['discount', 'DOUBLE', '折扣金额'],
         ['actual_payment', 'DOUBLE', '实际支付金额'],
         ['delivery_date / receive_date', 'VARCHAR', '发货/收货日期'],
         ['review_score', 'INT', '评价分数 (4-5)'],
         ['review_content', 'VARCHAR', '评价内容']],
        col_widths=[4, 4, 8]
    )

    add_heading(doc, '3.3 products — 产品表 (2,000行)', 2)
    add_styled_table(doc,
        ['字段', '类型', '说明'],
        [['product_id', 'VARCHAR(255)', '产品唯一标识'],
         ['product_name', 'VARCHAR(255)', '产品名称'],
         ['category', 'VARCHAR(255)', '品类 (15个品类)'],
         ['brand', 'VARCHAR(255)', '品牌 (113个品牌)'],
         ['price', 'DOUBLE', '产品标价'],
         ['sales_count', 'INT', '累计销量']],
        col_widths=[4, 4, 8]
    )

    add_heading(doc, '3.4 user_behaviors — 用户行为表 (30,000行)', 2)
    add_styled_table(doc,
        ['字段', '类型', '说明'],
        [['behavior_id', 'VARCHAR(255)', '行为唯一标识'],
         ['user_id', 'VARCHAR(255)', '用户ID'],
         ['product_id', 'VARCHAR(255)', '产品ID'],
         ['behavior_type', 'VARCHAR(255)', '行为类型 (浏览/点击/收藏/加购)'],
         ['behavior_time', 'DATETIME', '行为发生时间'],
         ['duration_seconds', 'INT', '行为持续秒数']],
        col_widths=[4, 4, 8]
    )

    add_heading(doc, '3.5 user_features — 用户特征表 (5,000行)', 2)
    add_styled_table(doc,
        ['字段', '类型', '说明'],
        [['user_id', 'VARCHAR(255)', '用户ID'],
         ['total_spent', 'DOUBLE', '累计消费金额'],
         ['order_count', 'DOUBLE', '订单总数'],
         ['avg_order_amount', 'DOUBLE', '平均订单金额'],
         ['browse/click/favorite/cart_count', 'DOUBLE', '各类行为次数'],
         ['days_since_last_order', 'INT', '距上次下单天数'],
         ['order_frequency', 'DOUBLE', '下单频次'],
         ['repurchase_indicator', 'INT', '复购标识'],
         ['purchase_intent', 'DOUBLE', '购买意向评分'],
         ['consumption_level', 'VARCHAR', '消费等级 (高/中/低)'],
         ['member_level_score', 'INT', '会员等级分']],
        col_widths=[4, 4, 8]
    )

    # === 4. 技术方案 ===
    add_heading(doc, '四、技术方案', 1)

    add_heading(doc, '4.1 技术栈', 2)
    add_styled_table(doc,
        ['层级', '技术选型', '用途'],
        [['数据存储', 'MySQL 8.0', '业务数据存储，utf8mb4编码'],
         ['数据提取', 'pymysql + pandas', '数据库连接、SQL查询、DataFrame构建'],
         ['数据处理', 'pandas + numpy', '数据清洗、聚合、特征工程'],
         ['机器学习', 'scikit-learn', 'K-Means聚类、StandardScaler标准化、PCA降维'],
         ['可视化', 'Plotly 6.x', '交互式图表 (18+种图表类型)'],
         ['BI前端', 'HTML5 + CSS Grid + JavaScript', 'Tab切换、响应式布局、CDN加载'],
         ['报告生成', 'python-docx', 'Word文档生成、表格样式、分页排版'],
         ['运行环境', 'Python 3.11+', 'Windows/Linux/Mac跨平台']],
        col_widths=[3, 4, 9]
    )

    add_heading(doc, '4.2 数据Pipeline架构', 2)
    add_body(doc, '数据流: MySQL(tmall_data) → pymysql连接 → pandas DataFrame → 特征工程(RFM+行为聚合) → '
              '分析计算(聚合/分组/队列) → 双输出(Plotly JSON→HTML / python-docx→Word)')
    add_body(doc, '聚类Pipeline: 特征选择(12维) → StandardScaler标准化 → K-Means(K=2~8评估) → '
              'Silhouette评分选K → PCA降维可视化 → 聚类画像统计 → 运营建议生成')

    add_heading(doc, '4.3 关键算法', 2)
    add_bullet(doc, 'RFM模型: Recency(最近消费距今天数) + Frequency(累计购买次数) + Monetary(累计消费金额)')
    add_bullet(doc, 'K-Means聚类: 肘部法则+轮廓系数确定最佳K值，业务场景选用K=5')
    add_bullet(doc, 'Cohort留存: 按用户首单周分组，计算每队列在第0-12周的留存率矩阵')
    add_bullet(doc, '生命周期分段: 基于R值(30天/90天阈值)分为活跃/沉默/流失三阶段')

    # === 5. 分析模块设计 ===
    add_heading(doc, '五、分析模块设计规格', 1)

    modules = [
        ('Tab 1 - 营收总览', '6 KPI + 6 图表',
         'GMV、实收、订单数、客单价、折扣率、复购率',
         '月度GMV/营收趋势(双轴柱状图)、日度营收+7日MA、月度订单量vs客单价、折扣率&ARPU趋势、支付方式分布(饼+柱)、订单状态分布(环形图)',
         '营收健康度评估、月度波动监测、支付渠道优化方向'),
        ('Tab 2 - 用户生命周期', '4 KPI + 6 图表',
         'DAU均值、粘性比(DAU/MAU)、平均复购间隔、用户平均LTV',
         'DAU趋势线、月度新增vs流失、生命周期阶段分布(环形图)、周度留存队列热力图、复购间隔分布直方图、LTV按订单数分布',
         '用户活跃度监控、流失预警、留存策略评估、获客成本参考'),
        ('Tab 3 - 产品分析', '4 图表 + 1 表格',
         '品类集中度、Top品牌贡献',
         '品类销售额排名(横向柱状图)、品牌Top10、品类均价对比、价格区间商品/销量双轴图、关联购买表',
         '品类资源分配、品牌合作决策、定价策略调整'),
        ('Tab 4 - 用户画像', '5 图表 + 1 交叉表',
         '性别比、年龄集中度、Top省份、消费等级分布',
         '性别/年龄/会员/地域/消费等级5维度图表、会员×消费交叉分析表',
         '用户特征洞察、区域运营策略、会员体系优化'),
        ('Tab 5 - 行为分析', '5 图表',
         '各阶段转化率、行为高峰时段',
         '转化漏斗(含%标注)、日行为量vs营收双轴、时段分布(小时级)、行为时长分布、用户活跃度分层',
         '转化瓶颈定位、推送时机优化、用户体验改善'),
        ('Tab 6 - 用户聚类', '4 KPI + 9 图表 + 5 建议',
         'K=5、轮廓系数、CH指数',
         '用户分层迁移桑基图(生命周期+RMF)、肘部法则、轮廓系数、PCA散点图、群体数量分布、RFM雷达图、群体性别/消费构成、特征表、迁移矩阵、5群体运营建议',
         '用户精细化分层、迁移趋势追踪、差异化营销策略、高价值用户维系'),
        ('Tab 7 - 评价服务', '4 KPI + 4 图表',
         '平均评分、平均配送天数、取消退款单数',
         '评分分布、月度评分趋势、物流时效分布、取消/退款支付方式分布',
         '服务质量管理、物流商评估、售后流程优化'),
    ]

    add_styled_table(doc,
        ['模块', '产出量', '核心指标', '可视化内容', '业务用途'],
        modules,
        col_widths=[2.5, 2, 2.5, 4.5, 4.5]
    )

    # === 6. 输出交付物 ===
    add_heading(doc, '六、输出交付物规格', 1)
    add_styled_table(doc,
        ['交付物', '格式', '规格', '说明'],
        [['统一BI报表', 'HTML', '~540KB, 8 Tab, 39+图表', '交互式BI看板，含首页导览+指标字典+7分析Tab+时间筛选滑块+用户分层迁移桑基图'],
         ['数据探索报告', 'Word (.docx)', '46KB, 10章', '表格化数据呈现+模块内嵌建议+优先行动事项表'],
         ['用户聚类脚本', 'Python (.py)', '562行', 'RFM特征工程+K-Means聚类+数据库写入+HTML报告'],
         ['BI报表脚本', 'Python (.py)', '~700行', '统一报表生成，含全部7Tab计算逻辑和HTML渲染'],
         ['PRD文档', 'Word (.docx)', '本文档', '项目背景、数据源、技术方案、模块规格、验收标准'],
         ['项目总结', 'Word (.docx)', '—', '设计思路、核心结论、业务价值、优化方向、交付物清单'],
         ['复现提示词', 'TXT', '—', '完整AI提示词，可按此复现整个项目'],
         ['数据库临时表', 'MySQL Table', 'claude_user_clusters', '5000用户聚类结果: user_id/cluster/cluster_name/RFM/PCA']],
        col_widths=[3, 2.5, 3, 7.5]
    )

    # === 7. 验收标准 ===
    add_heading(doc, '七、验收标准', 1)

    add_heading(doc, '7.1 数据准确性', 2)
    add_bullet(doc, 'BI报表KPI与Word报告数据100%一致（交叉验证通过）')
    add_bullet(doc, '用户聚类结果可复现（固定random_state=42，运行3次结果一致）')
    add_bullet(doc, '数据指标内部一致性：如总LTV × 用户数 ≈ 总营收')

    add_heading(doc, '7.2 功能完整性', 2)
    add_bullet(doc, 'HTML报表8个Tab全部可切换，图表正常渲染（≥39个Plotly图表，含2张桑基图）')
    add_bullet(doc, '时间筛选功能: 月份下拉选择器+应用/重置按钮，选择时间段后KPI自动更新、时序图表动态刷新')
    add_bullet(doc, '桑基图: Tab 6中生命周期阶段迁移图和RFM价值分层迁移图正常渲染，迁移矩阵表无错位')
    add_bullet(doc, 'Tab切换后Plotly图表自动resize，无错位或重叠')
    add_bullet(doc, 'Word文档无乱码，表格格式统一，分页正确')
    add_bullet(doc, 'Python脚本执行无报错（pandas Warning除外），可在同类环境一键运行')

    add_heading(doc, '7.3 用户体验', 2)
    add_bullet(doc, '响应式布局: 1100px以下3列变2列，768px以下全堆叠')
    add_bullet(doc, '首页指标字典覆盖12+核心指标的定义和公式')
    add_bullet(doc, '数据表格支持横向滚动，不超出容器宽度')
    add_bullet(doc, '每个Tab的分析结论位于顶部，用户切换即见洞察')

    # === 8. 项目依赖 ===
    add_heading(doc, '八、环境依赖与运行', 1)
    add_body(doc, '运行环境要求:')
    add_bullet(doc, 'Python 3.11+')
    add_bullet(doc, 'MySQL 8.0 (localhost, 数据库tmall_data, utf8mb4)')
    add_bullet(doc, 'Python包: pymysql, pandas, numpy, plotly, scikit-learn, python-docx')
    add_bullet(doc, '安装命令: pip install pymysql pandas numpy plotly scikit-learn python-docx')
    add_body(doc, '运行步骤:')
    add_bullet(doc, '1. 确保MySQL服务启动且tmall_data数据库可访问')
    add_bullet(doc, '2. 修改 user_auth_db.py 中的数据库密码(DB_CONFIG["password"])为实际密码')
    add_bullet(doc, '3. python unified_report.py → 生成统一BI报表HTML')
    add_bullet(doc, '4. python generate_exploration_report.py → 生成数据分析Word报告')
    add_bullet(doc, '5. python user_profiling.py → 运行聚类分析+生成HTML+写入数据库')

    return doc

if __name__ == '__main__':
    print("生成PRD文档...")
    doc = build_prd()
    output = r'D:\softdata\claude\test002\tmall_prd.docx'
    doc.save(output)
    print(f"PRD已生成: {output}")
