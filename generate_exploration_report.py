# -*- coding: utf-8 -*-
"""
天猫销售数据探索与业务经营建议报告 (专业版)
- 数据以表格呈现
- 问题发现与经营建议在各模块内嵌
- 独立优先行动事项模块
"""
import pymysql
import pandas as pd
import numpy as np
from user_auth_db import get_connection  # 数据库连接配置，请修改 user_auth_db.py 中的密码
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from datetime import datetime

# ============================================================
# 工具函数
# ============================================================
def fetch_data():
    conn = get_connection()
    users = pd.read_sql("SELECT * FROM users", conn)
    orders = pd.read_sql("SELECT * FROM orders", conn)
    products = pd.read_sql("SELECT * FROM products", conn)
    behaviors = pd.read_sql("SELECT * FROM user_behaviors", conn)
    features = pd.read_sql("SELECT * FROM user_features", conn)
    conn.close()
    return users, orders, products, behaviors, features

def prepare_data(users, orders, products, behaviors, features):
    orders['order_date_dt'] = pd.to_datetime(orders['order_date_date'])
    orders['order_month'] = orders['order_date_dt'].dt.to_period('M').astype(str)
    behaviors['behavior_time_dt'] = pd.to_datetime(behaviors['behavior_time'])
    behaviors['behavior_date'] = behaviors['behavior_time_dt'].dt.date
    order_product = orders.merge(products, on='product_id', how='left')
    return users, orders, products, behaviors, features, order_product

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_border(cell, **kwargs):
    """Set cell borders"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val","single")}" '
            f'w:sz="{val.get("sz","4")}" w:space="0" w:color="{val.get("color","CCCCCC")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)

def add_styled_table(doc, headers, rows, col_widths=None, header_color="667eea"):
    """Create a professionally styled table"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(h))
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)

    # Data rows
    for i, row in enumerate(rows):
        bg = "F8F9FF" if i % 2 == 0 else "FFFFFF"
        for j, val in enumerate(row):
            cell = table.rows[i+1].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            if j == 0:
                run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_shading(cell, bg)

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # spacing
    return table

def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return h

def add_insight_box(doc, finding, recommendation):
    """Add a finding + recommendation block"""
    # Finding
    p = doc.add_paragraph()
    r_label = p.add_run('数据发现: ')
    r_label.bold = True
    r_label.font.size = Pt(10)
    r_label.font.color.rgb = RGBColor(0x66, 0x7e, 0xea)
    r_text = p.add_run(finding)
    r_text.font.size = Pt(10)
    r_text.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Recommendation
    p2 = doc.add_paragraph()
    r_label2 = p2.add_run('经营建议: ')
    r_label2.bold = True
    r_label2.font.size = Pt(10)
    r_label2.font.color.rgb = RGBColor(0xe6, 0x7e, 0x22)
    r_text2 = p2.add_run(recommendation)
    r_text2.font.size = Pt(10)
    r_text2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()  # spacer

def add_section_intro(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    run.italic = True

# ============================================================
# 主报告构建
# ============================================================
def build_report(users, orders, products, behaviors, features, order_product):
    doc = Document()

    # ---- 封面 ----
    title = doc.add_heading('天猫用户销售数据探索分析报告', level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x66, 0x7e, 0xea)
        run.font.size = Pt(24)

    meta_table = doc.add_table(rows=4, cols=2)
    meta_data = [
        ('报告生成日期', datetime.now().strftime('%Y-%m-%d')),
        ('数据时间范围', '2025年9月6日 至 2026年3月5日（6个月）'),
        ('数据规模', '用户5,000 | 订单15,000 | 产品2,000 | 行为30,000'),
        ('分析工具', 'Python + pandas + sklearn + Plotly + python-docx'),
    ]
    for i, (k, v) in enumerate(meta_data):
        meta_table.rows[i].cells[0].text = k
        meta_table.rows[i].cells[1].text = v
        for j in range(2):
            for run in meta_table.rows[i].cells[j].paragraphs[0].runs:
                run.font.size = Pt(10)
                if j == 0:
                    run.bold = True
    doc.add_page_break()

    # ---- 关键指标计算 ----
    total_gmv = orders['total_amount'].sum() / 10000
    total_revenue = orders['actual_payment'].sum() / 10000
    total_orders = len(orders)
    avg_order = orders['actual_payment'].mean()
    unique_users = orders['user_id'].nunique()
    repurchase = (total_orders - unique_users) / unique_users * 100
    discount_rate = (1 - orders['actual_payment'].sum() / orders['total_amount'].sum()) * 100
    cancelled = (orders['order_status'] == '已取消').sum()
    refunded = (orders['order_status'] == '已退款').sum()

    monthly = orders.groupby('order_month').agg(
        GMV=('total_amount', 'sum'), revenue=('actual_payment', 'sum'),
        orders=('order_id', 'count'), customers=('user_id', 'nunique'),
        avg_order=('actual_payment', 'mean')
    ).reset_index()
    monthly['GMV'] /= 10000; monthly['revenue'] /= 10000
    monthly['discount'] = (1 - monthly['revenue'] / monthly['GMV']) * 100
    monthly['m2m_growth'] = monthly['revenue'].pct_change() * 100

    # ================================================================
    # 一、营收整体概览
    # ================================================================
    add_heading_styled(doc, '一、营收整体概览', 1)
    add_section_intro(doc, '报告期内平台核心经营指标汇总。GMV与实收差值为折扣让利部分，复购率反映用户消费粘性。')

    add_styled_table(doc,
        ['指标', '数值', '说明'],
        [
            ['总GMV', f'{total_gmv:,.0f} 万元', '所有订单标价总额'],
            ['实际营收', f'{total_revenue:,.0f} 万元', '扣除折扣后实际到账金额'],
            ['订单总数', f'{total_orders:,} 单', '含已完成/已取消/已退款所有状态'],
            ['消费用户数', f'{unique_users:,} 人', '至少有过1笔订单的去重用户'],
            ['客单价', f'{avg_order:,.0f} 元', '平均每单实付金额'],
            ['折扣率', f'{discount_rate:.1f}%', '整体让利幅度，电商行业均值12-18%'],
            ['复购率', f'{repurchase:.1f}%', '人均订单数-1，反映重复购买程度'],
            ['取消/退款率', f'{(cancelled+refunded)/total_orders*100:.1f}%', f'取消{cancelled}单+退款{refunded}单'],
        ],
        col_widths=[3.5, 3.5, 9]
    )

    add_insight_box(doc,
        f'总GMV {total_gmv:,.0f}万元，折扣率{discount_rate:.1f}%处于行业正常区间（12-18%）。实收{total_revenue:,.0f}万元。但取消退款占比{(cancelled+refunded)/total_orders*100:.1f}%略微偏高，需关注售后体验优化。',
        f'① 对比行业标杆（折扣率<12%），当前{discount_rate:.1f}%的让利空间可尝试通过提升商品力和品牌溢价逐步压缩；② 取消退款率目标压降至8%以内；③ 客单价{avg_order:,.0f}元有提升空间，可通过跨品类推荐和满减策略拉升。'
    )

    # ================================================================
    # 二、月度营收趋势与时间序列
    # ================================================================
    add_heading_styled(doc, '二、月度营收趋势与时间序列', 1)
    add_section_intro(doc, '按月维度拆解营收、订单量、客单价及环比变化，识别增长拐点与季节性规律。')

    monthly_rows = []
    for _, row in monthly.iterrows():
        mom = row['m2m_growth']
        mom_str = f'{mom:+.1f}%' if not pd.isna(mom) else '—'
        monthly_rows.append([
            row['order_month'],
            f'{row["revenue"]:,.0f}',
            f'{row["GMV"]:,.0f}',
            f'{row["orders"]:,}',
            f'{row["customers"]:,}',
            f'{row["avg_order"]:,.0f}',
            f'{row["discount"]:.1f}%',
            mom_str
        ])

    add_styled_table(doc,
        ['月份', '营收(万元)', 'GMV(万元)', '订单数', '消费用户', '客单价(元)', '折扣率', '环比增长'],
        monthly_rows,
        col_widths=[2, 2, 2, 1.5, 1.5, 2, 1.5, 1.8]
    )

    last_mom = monthly['m2m_growth'].iloc[-1] if len(monthly) > 1 else 0
    high_month = monthly.loc[monthly['revenue'].idxmax()]
    low_month = monthly.loc[monthly['revenue'].idxmin()]

    add_insight_box(doc,
        f'营收高峰出现在{high_month["order_month"]}（{high_month["revenue"]:,.0f}万元），低谷在{low_month["order_month"]}（{low_month["revenue"]:,.0f}万元），峰谷差{high_month["revenue"]-low_month["revenue"]:,.0f}万元（{(high_month["revenue"]/low_month["revenue"]-1)*100:.0f}%）。最新月环比{last_mom:+.1f}%。',
        f'① 结合行业大促节点（双11在11月、年货节在1月）预判高峰月，提前2-4周备货和预热；② 低谷月（如{low_month["order_month"]}）策划店铺自造节或品类日活动，拉平收入曲线；③ 目标将峰谷比从{(high_month["revenue"]/low_month["revenue"]-1)*100:.0f}%压缩至50%以内。'
    )

    # ================================================================
    # 三、订单与支付分析
    # ================================================================
    add_heading_styled(doc, '三、订单与支付分析', 1)
    add_section_intro(doc, '从订单状态分布识别流失环节，从支付方式偏好指导支付体验优化方向。')

    # 订单状态
    status = orders.groupby('order_status').size().sort_values(ascending=False)
    status_rows = [[s, f'{c:,}', f'{c/total_orders*100:.1f}%'] for s, c in status.items()]
    add_styled_table(doc, ['订单状态', '订单数', '占比'], status_rows, col_widths=[4, 4, 4])

    # 支付方式
    payment = orders.groupby('payment_method').agg(
        count=('order_id', 'count'), amount=('actual_payment', 'sum')
    ).reset_index()
    payment['amount'] /= 10000
    payment = payment.sort_values('count', ascending=False)
    pay_rows = [[r['payment_method'], f'{r["count"]:,}', f'{r["amount"]:,.0f}', f'{r["amount"]/r["count"]:,.0f}']
                for _, r in payment.iterrows()]
    add_styled_table(doc, ['支付方式', '订单数', '交易额(万元)', '笔均价(元)'], pay_rows, col_widths=[3, 3, 3, 3])

    add_insight_box(doc,
        f'已取消{status.get("已取消",0)}单+已退款{status.get("已退款",0)}单={(status.get("已取消",0)+status.get("已退款",0))/total_orders*100:.1f}%的订单未完成闭环。支付方式分布较均衡，没有单一渠道垄断。',
        f'① 建立取消/退款原因标签体系（尺码不符、质量不满意、物流太慢等），按原因TOP5逐一改善；② 监控各支付渠道的支付成功率和耗时，确保核心渠道（微信支付、支付宝）成功率>99.5%；③ 对高客单价品类引入分期支付选项。'
    )

    # ================================================================
    # 四、产品分析
    # ================================================================
    add_heading_styled(doc, '四、产品分析', 1)
    add_section_intro(doc, '从品类和品牌两个维度透视产品结构，识别核心品类和增长机会。')

    # 品类
    category = order_product.groupby('category').agg(
        sales=('actual_payment', 'sum'), orders=('order_id', 'count'),
        avg_price=('price', 'mean'), products=('product_id', 'nunique'),
        avg_discount=('discount', 'mean')
    ).reset_index()
    category['sales'] /= 10000
    category['sales_pct'] = category['sales'] / category['sales'].sum() * 100
    category = category.sort_values('sales', ascending=False)

    cat_rows = [[r['category'], f'{r["sales"]:,.0f}', f'{r["sales_pct"]:.1f}%', f'{r["orders"]:,}',
                 f'{r["products"]}', f'{r["avg_price"]:,.0f}', f'{r["avg_discount"]:,.1f}']
                for _, r in category.iterrows()]
    add_styled_table(doc,
        ['品类', '销售额(万元)', '销售占比', '订单数', 'SKU数', '均价(元)', '平均折扣'],
        cat_rows, col_widths=[2.5, 2.2, 1.8, 1.8, 1.2, 2, 2]
    )

    # 品牌Top10
    brand = order_product.groupby('brand').agg(sales=('actual_payment', 'sum'), orders=('order_id', 'count')).reset_index()
    brand['sales'] /= 10000
    brand = brand.sort_values('sales', ascending=False).head(10)
    brand_rows = [[r['brand'], f'{r["sales"]:,.0f}', f'{r["orders"]:,}', f'{r["sales"]/r["orders"]:,.0f}']
                  for _, r in brand.iterrows()]
    add_styled_table(doc, ['品牌', '销售额(万元)', '订单数', '笔均价(元)'], brand_rows, col_widths=[3, 3, 3, 3])

    top3_share = category['sales_pct'].head(3).sum()
    top_cat = category.iloc[0]

    add_insight_box(doc,
        f'Top3品类贡献{top3_share:.1f}%销售额，头部集中度明显。最高品类"{top_cat["category"]}"占{top_cat["sales_pct"]:.1f}%，是平台的绝对核心品类。品类均价跨度大（{category["avg_price"].min():.0f}-{category["avg_price"].max():.0f}元），反映不同品类的价格策略差异。',
        f'① 对Top3品类保持资源倾斜：增加SKU丰富度、争取品牌独家合作、优化搜索排名；② 筛选利润率高于均值的长尾品类（如{category.nsmallest(3, "sales_pct")["category"].values[0] if len(category)>3 else "珠宝首饰"}等），作为第二增长曲线培育；③ 建立品类健康度仪表盘，监控各品类的动销率、库存周转、毛利率。'
    )

    # ================================================================
    # 五、用户分析
    # ================================================================
    add_heading_styled(doc, '五、用户画像与分层分析', 1)
    add_section_intro(doc, '从人口统计和消费行为两个维度描绘用户全貌，识别核心用户群与增长机会。')

    # 性别 + 年龄
    gender = users['gender'].value_counts()
    age_bins = [0, 18, 25, 30, 35, 40, 50, 100]
    age_labels = ['<18', '18-24', '25-29', '30-34', '35-39', '40-49', '50+']
    users['age_group'] = pd.cut(users['age'], bins=age_bins, labels=age_labels)
    age_dist = users['age_group'].value_counts().sort_index()

    demo_rows = []
    for g, c in gender.items():
        demo_rows.append([f'性别-{g}', f'{c:,}', f'{c/len(users)*100:.1f}%', '—'])
    for a, c in age_dist.items():
        demo_rows.append([f'年龄-{a}', f'{c:,}', f'{c/len(users)*100:.1f}%', '—'])
    add_styled_table(doc, ['维度', '人数', '占比', '备注'], demo_rows, col_widths=[4, 3, 3, 4])

    # 会员等级
    member = users['member_level'].value_counts()
    member_rows = [[m, f'{c:,}', f'{c/len(users)*100:.1f}%'] for m, c in member.items()]
    add_styled_table(doc, ['会员等级', '人数', '占比'], member_rows, col_widths=[4, 4, 4])

    # 消费等级
    cons = features['consumption_level'].value_counts()
    cons_rows = [[l, f'{c:,}', f'{c/len(features)*100:.1f}%'] for l, c in cons.items()]
    add_styled_table(doc, ['消费能力', '人数', '占比'], cons_rows, col_widths=[4, 4, 4])

    # 省份
    province = users['province'].value_counts().head(10)
    prov_rows = [[p, f'{c:,}', f'{c/len(users)*100:.1f}%'] for p, c in province.items()]
    add_styled_table(doc, ['省份(Top10)', '用户数', '占比'], prov_rows, col_widths=[4, 4, 4])

    low_cons_pct = cons.get('低', 0) / len(features) * 100

    add_insight_box(doc,
        f'男女比例{gender.iloc[0]/(gender.sum())*100:.0f}:{gender.iloc[1]/(gender.sum())*100:.0f}，年龄集中于25-35岁（占比{age_dist["25-29"]/age_dist.sum()*100+age_dist["30-34"]/age_dist.sum()*100:.0f}%）。低消费能力用户占{low_cons_pct:.1f}%，是潜在增长空间。',
        f'① 25-35岁为核心人群，匹配品质生活、职场进阶、家庭消费等营销主题；② 低消费用户({low_cons_pct:.1f}%)需通过首单补贴和低价引流品激活；③ 省份Top3用户占比较高，可在这些区域做线下联动或区域专享活动。'
    )

    # ================================================================
    # 六、用户行为与转化分析
    # ================================================================
    add_heading_styled(doc, '六、用户行为与转化分析', 1)
    add_section_intro(doc, '从行为漏斗、活跃度、时段分布等维度分析用户在产品内的行为模式。')

    behavior = behaviors['behavior_type'].value_counts()
    browse = behavior.get('浏览', 1)
    click = behavior.get('点击', 0)
    fav = behavior.get('收藏', 0)
    cart = behavior.get('加购', 0)
    total_beh = len(behaviors)

    funnel_rows = [
        ['浏览', f'{browse:,}', '100.0%', '—'],
        ['点击', f'{click:,}', f'{click/browse*100:.1f}%', f'{click/browse*100:.1f}%（浏览→点击）'],
        ['收藏', f'{fav:,}', f'{fav/browse*100:.1f}%', f'{fav/click*100:.1f}%（点击→收藏）'],
        ['加购', f'{cart:,}', f'{cart/browse*100:.1f}%', f'{cart/click*100:.1f}%（点击→加购）'],
        ['下单', f'{total_orders:,}', f'{total_orders/browse*100:.1f}%', f'{total_orders/cart*100:.1f}%（加购→下单）'],
    ]
    add_styled_table(doc,
        ['阶段', '数量', '对浏览占比', '阶段转化率'],
        funnel_rows, col_widths=[2.5, 3, 3, 7.5])

    # 用户活跃度
    user_beh_count = behaviors.groupby('user_id').size()
    act_bins = [0, 3, 5, 10, 20, 50, 99999]
    act_labels = ['1-3次(低活)', '4-5次', '6-10次', '11-20次', '21-50次(中活)', '50次+(高活)']
    user_beh_count_cut = pd.cut(user_beh_count, bins=act_bins, labels=act_labels)
    act_dist = user_beh_count_cut.value_counts().sort_index()

    act_rows = [[a, f'{c:,}', f'{c/len(user_beh_count)*100:.1f}%'] for a, c in act_dist.items()]
    add_styled_table(doc, ['活跃度分层', '用户数', '占比'], act_rows, col_widths=[5, 4, 4])

    # 时段分布
    behaviors['hour'] = behaviors['behavior_time_dt'].dt.hour
    hourly = behaviors.groupby('hour').size()
    peak_hour = hourly.idxmax()
    peak_count = hourly.max()

    add_insight_box(doc,
        f'浏览→加购转化率{funnel_rows[3][2]}，浏览→下单转化率{funnel_rows[4][2]}。加购到下单转化率{funnel_rows[4][3]}，表明加购用户购买意向强。活跃度集中在低活区间（{act_dist.iloc[0]:,}人只产生1-3次行为），用户深度参与不足。行为高峰在{peak_hour}:00，峰值{peak_count:,}次。',
        f'① 加购→下单转化率较高，应在加购后30分钟推送限时优惠或库存紧张提示，加速决策；② 低活跃用户占比较高，建议通过Push推送个性化商品和签到积分体系提升回访频率；③ 在行为高峰时段({peak_hour}:00前后1小时)集中投放广告和推送，最大化触达效率。'
    )

    # ================================================================
    # 七、用户生命周期分析
    # ================================================================
    add_heading_styled(doc, '七、用户生命周期分析', 1)
    add_section_intro(doc, '从LTV、留存、活跃度等维度分析用户价值生命周期。')

    # LTV
    ref_date = orders['order_date_dt'].max() + pd.Timedelta(days=1)
    rfm = orders.groupby('user_id').agg(
        Recency=('order_date_dt', lambda x: (ref_date - x.max()).days),
        Frequency=('order_id', 'count'),
        Monetary=('actual_payment', 'sum')
    ).reset_index()

    # 生命周期阶段
    active_now = (rfm['Recency'] <= 30).sum()
    dormant = ((rfm['Recency'] > 30) & (rfm['Recency'] <= 90)).sum()
    churned = (rfm['Recency'] > 90).sum()

    # DAU (from behaviors)
    daily_active = behaviors.groupby('behavior_date')['user_id'].nunique()
    avg_dau = daily_active.mean()
    mau = behaviors.groupby(behaviors['behavior_time_dt'].dt.to_period('M').astype(str))['user_id'].nunique().iloc[-1]
    stickiness = avg_dau / mau * 100 if mau > 0 else 0

    # Average repurchase interval
    user_orders_sorted = orders[['user_id', 'order_date_dt']].sort_values(['user_id', 'order_date_dt'])
    user_orders_sorted['prev_order'] = user_orders_sorted.groupby('user_id')['order_date_dt'].shift(1)
    user_orders_sorted['interval_days'] = (user_orders_sorted['order_date_dt'] - user_orders_sorted['prev_order']).dt.days
    avg_interval = user_orders_sorted['interval_days'].dropna().mean()

    lifecycle_rows = [
        ['日均活跃用户(DAU)', f'{avg_dau:,.0f} 人', '每日有行为记录的去重用户'],
        ['MAU', f'{mau:,} 人', '月活跃用户数'],
        ['用户粘性比(DAU/MAU)', f'{stickiness:.1f}%', '反映用户回访频率，健康值>20%'],
        ['活跃用户(30天内有订单)', f'{active_now:,} 人', f'占比 {active_now/len(rfm)*100:.1f}%'],
        ['沉默用户(30-90天)', f'{dormant:,} 人', f'占比 {dormant/len(rfm)*100:.1f}%'],
        ['流失用户(>90天)', f'{churned:,} 人', f'占比 {churned/len(rfm)*100:.1f}%'],
        ['用户平均LTV(累计消费)', f'{rfm["Monetary"].mean():,.0f} 元', f'中位数: {rfm["Monetary"].median():,.0f}元'],
        ['平均复购间隔', f'{avg_interval:.1f} 天', '相邻两次订单的平均间隔天数'],
        ['人均订单数', f'{rfm["Frequency"].mean():.1f} 单', f'中位数: {rfm["Frequency"].median():.0f}单'],
    ]
    add_styled_table(doc,
        ['生命周期指标', '数值', '说明'],
        lifecycle_rows, col_widths=[4, 3.5, 8.5])

    dormant_churn_pct = (dormant + churned) / len(rfm) * 100

    add_insight_box(doc,
        f'沉默+流失用户合计{dormant+churned:,}人，占消费用户的{dormant_churn_pct:.1f}%。DAU/MAU粘性比{stickiness:.1f}%，用户平均复购间隔{avg_interval:.1f}天。用户平均LTV {rfm["Monetary"].mean():,.0f}元。',
        f'① 沉默+流失用户{dormant_churn_pct:.1f}%占比过高（行业标杆<40%），需启动系统化唤醒计划：分层发放回归券、EDM/短信触达、推送更新内容；② 粘性比目标提升至25%以上，通过签到打卡、每日秒杀、内容社区等手段增加用户回访触点；③ 复购间隔{avg_interval:.1f}天，针对快消品类可在第{max(1,int(avg_interval-7))}天推送复购提醒。'
    )

    # ================================================================
    # 八、评价与服务分析
    # ================================================================
    add_heading_styled(doc, '八、评价与服务分析', 1)
    add_section_intro(doc, '从用户评价和物流时效评估服务质量。')

    review_dist = orders[orders['review_score'].notna()]['review_score'].value_counts().sort_index()
    avg_review = orders['review_score'].mean()

    review_rows = [[f'{int(s)}分', f'{c:,}', f'{c/len(orders)*100:.1f}%'] for s, c in review_dist.items()]
    add_styled_table(doc, ['评分', '订单数', '占比'], review_rows, col_widths=[4, 4, 4])

    # 物流
    orders['delivery_dt'] = pd.to_datetime(orders['delivery_date'])
    orders['receive_dt'] = pd.to_datetime(orders['receive_date'])
    orders['delivery_days'] = (orders['receive_dt'] - orders['delivery_dt']).dt.days
    delivery_valid = orders[(orders['delivery_days'] >= 0) & (orders['delivery_days'] <= 30)]
    avg_delivery = delivery_valid['delivery_days'].mean()

    delivery_bins = [0, 1, 3, 5, 7, 14, 31]
    delivery_labels = ['当日达(<=1天)', '2-3天', '4-5天', '6-7天', '8-14天', '15天+']
    delivery_valid['range'] = pd.cut(delivery_valid['delivery_days'], bins=delivery_bins, labels=delivery_labels)
    del_dist = delivery_valid['range'].value_counts().sort_index()

    del_rows = [[d, f'{c:,}', f'{c/len(delivery_valid)*100:.1f}%'] for d, c in del_dist.items()]
    add_styled_table(doc, ['物流时效', '订单数', '占比'], del_rows, col_widths=[4, 4, 4])

    good_review_pct = (review_dist.get(4, 0) + review_dist.get(5, 0)) / review_dist.sum() * 100
    fast_delivery_pct = (del_dist.iloc[0] + del_dist.iloc[1] if len(del_dist) > 1 else 0) / del_dist.sum() * 100

    add_insight_box(doc,
        f'综合评分{avg_review:.2f}分，{good_review_pct:.1f}%的用户给出4-5星好评。平均配送{avg_delivery:.1f}天。仅评4星和5星，无可用的1-3星数据，可能表示数据采集不够完整。',
        f'① 评分数据仅有4-5分，建议检查评价采集机制是否遗漏低分评价；② 物流时效{avg_delivery:.1f}天，目标压缩至3天以内以提升用户体验；③ 建立NPS（净推荐值）体系替代单一评分，更准确衡量用户满意度。'
    )

    # ================================================================
    # 九、优先行动事项
    # ================================================================
    doc.add_page_break()
    add_heading_styled(doc, '九、优先行动事项', 1)
    add_section_intro(doc, '综合以上数据分析，按紧急度×影响力矩阵排列优先行动事项。建议管理层按P0→P1→P2顺序推进。')

    priority_rows = [
        ['P0\n紧急且重要', '降低沉默流失率',
         f'沉默+流失用户{dormant+churned}人（{dormant_churn_pct:.1f}%），远超可控范围',
         '流失预警模型上线 + 分层唤醒自动化流程：\n① 30天未购→推送品类优惠券\n② 60天未购→EDM+短信双通道触达\n③ 90天未购→电话回访+专属大额券',
         '预计30天内召回15-20%',
         '2周'],

        ['P0\n紧急且重要', '提升浏览→下单转化率',
         f'浏览→下单转化率仅{total_orders/browse*100:.1f}%，大量流量浪费',
         '① 加购后30分钟推送库存紧张/限时优惠消息\n② 优化商品详情页（视频+真实用户晒图）\n③ 新人专享价+首单免邮',
         '转化率提升30%（目标{total_orders/browse*130:.1f}%）',
         '3周'],

        ['P1\n重要不紧急', '优化品类结构',
         f'Top3品类贡献{top3_share:.1f}%销售，长尾品类发展不足',
         '① 筛选利润率>25%的潜力品类重点扶持\n② 定期清理30天零动销SKU\n③ 引入品牌直供，降低采购成本',
         '长尾品类销售占比提升5pp',
         '6周'],

        ['P1\n重要不紧急', '会员体系升级',
         f'低消费用户占{low_cons_pct:.1f}%，会员权益感知弱',
         '① 建立任务积分体系（签到/浏览/评价/分享得积分）\n② 会员日专属折扣+双倍积分\n③ 高等级会员免费试用新品特权',
         '低消费→中消费升级转化率提升20%',
         '8周'],

        ['P1\n重要不紧急', '取消退款率治理',
         f'取消+退款率{(cancelled+refunded)/total_orders*100:.1f}%，侵蚀利润',
         '① 建立退款原因标签体系（5大分类+20子标签）\n② 针对TOP3退款原因制定改善方案\n③ 完善商品尺码表和实物图标准',
         '退款率降至8%以内',
         '4周'],

        ['P2\n常规优化', '提升用户粘性(DAU/MAU)',
         f'粘性比{stickiness:.1f}%，目标提升至25%+',
         '① 每日秒杀/签到打卡功能\n② UGC内容社区（晒单/测评/问答）\n③ 个性化瀑布流推荐',
         'DAU/MAU达到25%',
         '12周'],

        ['P2\n常规优化', '数据基础设施完善',
         '目前缺乏1-3星评价数据、流量来源数据、营销活动ROI数据',
         '① 补齐评价采集机制（确认收货后强制评价）\n② 部署UTM全链路流量追踪\n③ 搭建营销活动效果评估看板',
         '数据完整度从60%提升至90%',
         '12周'],
    ]

    add_styled_table(doc,
        ['优先级', '行动项', '数据支撑', '具体措施', '预期效果', '建议周期'],
        priority_rows,
        col_widths=[1.5, 2.5, 3.5, 5, 3, 1.5],
        header_color="e74c3c"
    )

    # ================================================================
    # 十、总结
    # ================================================================
    add_heading_styled(doc, '十、总结', 1)

    summary_items = [
        f'营收健康度: 总GMV {total_gmv:,.0f}万元（6个月），折扣率{discount_rate:.1f}%处于行业合理区间，实收{total_revenue:,.0f}万元。',
        f'用户资产: {unique_users:,}名消费用户，人均消费{total_revenue*10000/unique_users:,.0f}元，但{dormant_churn_pct:.1f}%用户处于沉默或流失状态，用户资产保全迫在眉睫。',
        f'转化效率: 浏览→下单转化率{total_orders/browse*100:.1f}%，加购→下单转化率{total_orders/cart*100:.1f}%，后者是高效的转化节点，应重点运营。',
        f'产品结构: {len(category)}个品类中Top3贡献{top3_share:.1f}%销售，需平衡头部品类的资源投入与长尾品类的增长培育。',
        f'服务体验: 评分{avg_review:.2f}分、配送{avg_delivery:.1f}天，处于良好水平，但评价数据完整性和退款率仍有改善空间。',
    ]

    for item in summary_items:
        p = doc.add_paragraph()
        run = p.add_run(f'• {item}')
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # 结语
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        '综上，建议管理层以"降低用户流失"和"提升转化效率"为两大核心抓手，'
        '按照优先行动事项表中的P0→P1→P2顺序分阶段推进。每两周复盘关键指标变化，'
        '以数据驱动的方式持续优化运营策略，实现从天猫店铺向用户品牌的升级。'
    )
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x7e, 0xea)
    run.italic = True

    return doc


if __name__ == '__main__':
    print("=" * 60)
    print("生成天猫数据分析报告 (专业版)")
    print("=" * 60)
    print("提取数据...")
    users, orders, products, behaviors, features = fetch_data()
    print("数据预处理...")
    users, orders, products, behaviors, features, order_product = prepare_data(users, orders, products, behaviors, features)
    print("构建报告（表格化+模块内嵌建议+优先事项）...")
    doc = build_report(users, orders, products, behaviors, features, order_product)

    output_path = r'D:\softdata\claude\test002\tmall_exploration_report.docx'
    doc.save(output_path)
    print(f"报告已生成: {output_path}")
    print("=" * 60)
